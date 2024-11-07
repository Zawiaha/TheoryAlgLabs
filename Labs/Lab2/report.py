from docx import Document
import openpyxl

def save_to_docx(data, filename='report.docx'):
    doc = Document()
    doc.add_heading('Trip Report', level=1)
    
    for mode, info in data.items():
        doc.add_paragraph(f'Тип транспорта: {mode}')
        doc.add_paragraph(f'Время пути: {info["time"]:.2f} hours')
        doc.add_paragraph(f'Стоимость: ${info["cost"]:.2f}')
    
    doc.save(filename)

def save_to_xls(data, filename='report.xlsx'):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Trip Report"
    
    ws.append(["Transport", "Time (hours)", "Cost ($)"])
    
    for mode, info in data.items():
        ws.append([mode, info["time"], info["cost"]])
    
    wb.save(filename)